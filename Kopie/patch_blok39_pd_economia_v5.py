# -*- coding: utf-8 -*-
"""
Nadpisuje kolumny 4–7 (Źródło, Opis logiki, Moment, Znaczenie) dla 39 wierszy
między neu_quantityfromexcel a «Produkt przetargu» w opis_pol_przyciskow_akcji_v5.docx.

Uruchomienie (Kopie): python patch_blok39_pd_economia_v5.py
Potem: python _fill_formularz_forward_blok_v5.py (jeśli kolumna Formularz była pusta w kontynuacji bloku)
Potem: python reformat_poc_pipeline_polski_docx.py "opis_pol_przyciskow_akcji_v5.docx" --backup
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from blok39_pd_economia_opisy import MOMENT_BLOK, OPIS_LOGIKI, ZNACZENIE, ZRODLO_BLOK
from docx_cell_format import cell_paragraph_joined_text, clear_cell

DOC = Path(__file__).resolve().parent / "opis_pol_przyciskow_akcji_v5.docx"


def _set_cell_lines(cell, text: str) -> None:
    clear_cell(cell)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.add_run(line)


def _find_anchor(table) -> int:
    for i, row in enumerate(table.rows):
        if cell_paragraph_joined_text(row.cells[2]).strip() == "neu_quantityfromexcel":
            return i
    raise RuntimeError("Brak neu_quantityfromexcel.")


def _find_produkt_przetargu(table) -> int:
    for i, row in enumerate(table.rows):
        if cell_paragraph_joined_text(row.cells[0]).strip() == "Produkt przetargu":
            return i
    raise RuntimeError('Brak wiersza Formularz == "Produkt przetargu".')


def main() -> None:
    doc = Document(str(DOC))
    table = doc.tables[0]
    a = _find_anchor(table)
    b = _find_produkt_przetargu(table)
    n = b - a - 1
    if n != 39:
        raise RuntimeError(f"Oczekiwano 39 wierszy między kotwicami, jest {n}")

    patched = 0
    for j in range(39):
        row = table.rows[a + 1 + j]
        logic = cell_paragraph_joined_text(row.cells[2]).strip()
        if logic not in OPIS_LOGIKI:
            raise KeyError(f"Brak OPIS_LOGIKI dla {logic} (wiersz {a+1+j})")
        _set_cell_lines(row.cells[4], ZRODLO_BLOK)
        _set_cell_lines(row.cells[5], OPIS_LOGIKI[logic])
        _set_cell_lines(row.cells[6], MOMENT_BLOK)
        _set_cell_lines(row.cells[7], ZNACZENIE[logic])
        patched += 1

    doc.save(str(DOC))
    print("Zapisano:", DOC)
    print("Zaktualizowano wierszy:", patched)


if __name__ == "__main__":
    main()
